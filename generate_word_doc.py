import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_document():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("PBR VITS CAMPUS COMPANION")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138) # Deep Blue

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Smart Campus Management, AI Face Recognition Biometrics & Scheduling Ecosystem\nPBR Visvodaya Institute of Technology & Science (Kavali, AP)")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph() # Spacer

    # Section 1: Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Project Purpose", level=1)
    h1.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    p1 = doc.add_paragraph(
        "PBR VITS Campus Companion is an enterprise-grade web application engineered to modernize daily academic "
        "and administrative workflows at PBR Visvodaya Institute of Technology and Science. The system replaces manual, "
        "paper-based, and proxy-vulnerable attendance routines with a high-accuracy, AI-powered Face Recognition System (FRS), "
        "coupled with a real-time academic timetable engine, dynamic faculty profiling, attendance analytics with exam eligibility predictors, "
        "and an integrated interactive AI study companion."
    )
    p1.runs[0].font.size = Pt(11)

    # Section 2: Technology Stack & Applications
    h2 = doc.add_heading("2. Comprehensive Technology Stack", level=1)
    h2.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph("The application leverages modern, high-performance web and artificial intelligence technologies across all layers:")

    # Tech Stack Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Layer / Component", "Technology / Framework", "Role & Functionality"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A")
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10.5)

    tech_data = [
        ("Frontend Framework", "React 19, TypeScript, Vite", "Ultra-fast Single Page Application (SPA), component architecture, strict type checking, and modular routing."),
        ("User Interface & Styling", "Tailwind CSS, Lucide React", "Responsive mobile-first layout, custom glassmorphism components, animated progress bars, and accessible iconography."),
        ("Artificial Intelligence / ML", "@vladmandic/face-api, TensorFlow.js", "Client-side neural networks (SSD MobileNet V1, 68-point facial landmark detector, 128-dimensional vector descriptor extractor)."),
        ("Camera & Media Streaming", "WebRTC MediaStream API", "Real-time webcam video feed capture, orientation alignment, multi-frame variant sampling, and live facial frame tracking."),
        ("Backend Web API", "FastAPI (Python 3.11+), Uvicorn", "High-throughput asynchronous REST API, Pydantic input schemas, Euclidean vector similarity computation, and security headers."),
        ("Database & ORM Layer", "PostgreSQL (Supabase) + SQLAlchemy", "Relational persistence, database connection pooling, auto-schema migrations, and foreign-key integrity."),
        ("Cloud Hosting & CI/CD", "Vercel + Render + GitHub Actions", "Frontend deployed on Vercel Edge Network, Backend hosted on Render cloud, synchronized via automated GitHub CI/CD.")
    ]

    for layer, tech, role in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[1].text = tech
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[2].text = role
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_paragraph() # Spacer

    # Section 3: Core Functional Modules
    h3 = doc.add_heading("3. Core Functional Modules & Operational Workflows", level=1)
    h3.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    # 3.1 Face Recognition System
    h3_1 = doc.add_heading("3.1. AI Face Recognition Smart Attendance (FRS)", level=2)
    h3_1.runs[0].font.color.rgb = RGBColor(14, 116, 144)

    frs_points = [
        ("Multi-Variant Biometric Enrollment: ", "During registration, the student's face is scanned across varying lighting and angle conditions, producing a set of normalized 128-dimensional Euclidean feature vectors stored securely in the database."),
        ("Strict 25-Minute Time Window: ", "Attendance for any scheduled period opens 10 minutes prior to class start time and closes 15 minutes after start time (25-minute active window). Scans attempted outside this interval are blocked as 'Window Closed' to prevent attendance spoofing."),
        ("Euclidean Distance & Confidence Match: ", "When a student verifies their face, the live vector is matched against their enrolled embedding using Euclidean distance algorithms (minimum distance across variants with confidence scoring >= 75%)."),
        ("Period-Isolated Presence: ", "Attendance is marked strictly for that individual class period (e.g. Period 3 - Expert Systems), ensuring accurate session-by-session presence tracking.")
    ]
    for bold_txt, norm_txt in frs_points:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_txt)
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r2 = p.add_run(norm_txt)
        r2.font.size = Pt(10.5)

    # 3.2 Timetable Engine
    h3_2 = doc.add_heading("3.2. Academic Timetable & Smart Scheduling Matrix", level=2)
    h3_2.runs[0].font.color.rgb = RGBColor(14, 116, 144)

    tt_points = [
        ("Comprehensive Multi-Department Support: ", "Pre-configured for 6 engineering disciplines: Computer Science and Engineering (CSE), CSE AI, CSE AIML, Electronics and Communication Engineering (ECE), Electrical and Electronics Engineering (EEE), and Civil Engineering."),
        ("8-Semester Structure: ", "Full weekly schedules (Periods 1 to 4+ from Monday to Friday) for all semesters from 1-1 up to 4-2."),
        ("Faculty Self-Service Scheduling: ", "Faculty members can add new periods, schedule afternoon/lab sessions (e.g. Period 5), and reassign classrooms with built-in conflict detection to prevent double-booking teachers or lecture halls.")
    ]
    for bold_txt, norm_txt in tt_points:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_txt)
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r2 = p.add_run(norm_txt)
        r2.font.size = Pt(10.5)

    # 3.3 Faculty Directory
    h3_3 = doc.add_heading("3.3. Faculty Management & Academic Profiling", level=2)
    h3_3.runs[0].font.color.rgb = RGBColor(14, 116, 144)
    p_fac = doc.add_paragraph(
        "A directory containing 53+ official PBR VITS faculty members across all departments with verified credentials, "
        "doctoral degrees (Ph.D.), dates of joining, university backgrounds, designations (Professors, Associate Professors, Assistant Professors), "
        "and assigned subjects/semesters."
    )
    p_fac.runs[0].font.size = Pt(10.5)

    # 3.4 Attendance Analytics & Predictor
    h3_4 = doc.add_heading("3.4. Attendance Analytics & Exam Eligibility Predictor", level=2)
    h3_4.runs[0].font.color.rgb = RGBColor(14, 116, 144)
    p_calc = doc.add_paragraph(
        "Dynamic subject-wise attendance monitors calculate current presence percentages against the university's mandatory 75% threshold. "
        "The built-in interactive simulator enables students to calculate exactly how many upcoming lectures they must attend to achieve or maintain exam eligibility."
    )
    p_calc.runs[0].font.size = Pt(10.5)

    # 3.5 AI Study Companion
    h3_5 = doc.add_heading("3.5. Interactive AI Academic Study Companion", level=2)
    h3_5.runs[0].font.color.rgb = RGBColor(14, 116, 144)
    p_ai = doc.add_paragraph(
        "Integrated AI chatbot trained on university curriculum topics, providing instant answers to academic questions, "
        "study notes summaries, laboratory guides, and exam preparation assistance."
    )
    p_ai.runs[0].font.size = Pt(10.5)

    # Section 4: Security & Anti-Spoofing
    h4 = doc.add_heading("4. Security, Anti-Spoofing & Data Privacy", level=1)
    h4.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    sec_points = [
        ("No Raw Image Storage: ", "The system does not persist raw biometric facial photographs; instead, only normalized 128-float mathematical embeddings are stored, ensuring privacy compliance."),
        ("Multi-Variant Lighting Tolerance: ", "Robust embedding distance thresholds eliminate false rejections caused by varying campus lighting conditions."),
        ("Role-Based Access Control (RBAC): ", "Strict access boundaries between Students, Faculty, Admins, and HODs enforced via backend HTTP headers and database token validation.")
    ]
    for bold_txt, norm_txt in sec_points:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_txt)
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        r2 = p.add_run(norm_txt)
        r2.font.size = Pt(10.5)

    # Section 5: Project Access & Deployment Links
    h5 = doc.add_heading("5. Production Deployment Links", level=1)
    h5.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    p_links = doc.add_paragraph()
    p_links.add_run("• Live Web Application (Vercel): ").font.bold = True
    p_links.add_run("https://pbr-vits-companion.vercel.app\n")
    p_links.add_run("• Backend REST API (Render): ").font.bold = True
    p_links.add_run("https://campus-companion-prototype.onrender.com\n")
    p_links.add_run("• Source Code Repository (GitHub): ").font.bold = True
    p_links.add_run("https://github.com/Indorusky/pbr-vits-companion\n")

    # Save document
    output_path = r"c:\a\DOC-20260822-WA0007\campus_companion\campus_companion_prototyoe (2)\campus_companion_prototyoe\PBR_VITS_Campus_Companion_Project_Documentation.docx"
    doc.save(output_path)
    print(f"Word document successfully created at: {output_path}")

if __name__ == "__main__":
    create_document()
